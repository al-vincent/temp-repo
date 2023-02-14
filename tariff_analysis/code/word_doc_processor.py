"""
This module ingests tables contained within two documents: 
1. "Authorised-Use-Eligible-goods-and-authorised-uses-Final-20210719__1_.docx"
2. "Authorised-Use-Eligible-goods-and-rates-Final-20210719.docx"

In both cases, the module extracts the tables required, applies some
processing, and outputs the cleaned data to a CSV file. 

There is considerable overlap between the two documents. For example,
the Rates document contains all the descriptive and usage information
that's contained in the Usage document, for the comm codes within it. 
However, some points to note:
- There are a lot of commodity codes in the Uses doc that don't appear
in the Rates doc.
- There are a small number of comm codes in the Rates doc that don't
appear in the Uses doc (approx 5).

Therefore, it is useful to process both documents separately, 
extracting rates from one and usage / commodity descriptions from the
other.
"""

#======================================================================================
# Import libraries
#======================================================================================
import argparse
import logging
import logging.config
import os
from datetime import datetime

from docx import Document
import numpy as np
import pandas as pd

import config as c
from base import WordDocProcessor
import utils as ut
import log_config as lc

#======================================================================================
# Setup Logger
#======================================================================================
logging.config.dictConfig(lc.LOG_CONFIG)
logger = logging.getLogger(__name__)


#======================================================================================
# Classes to process the AU Reference Documents
#======================================================================================

# -------------
# AU Rates Doc
# -------------
class AuRatesDocIngest(WordDocProcessor):
    def __init__(self, input_file, output_file, show_progress=True):
        super().__init__(input_file)
        self.output_file = output_file
        self.show_progress = show_progress
    
    def process_all_tables(self):        
        au_rates_tables = []
        idx = 0
        if self.show_progress:
            print(f"Number of AU Rates tables in document: {len(self.doc.tables)}")
        for table in self.doc.tables:
            if len(table.columns) == 3:
                idx += 1
                tbl = self.docx_to_list(table, first_row_as_hdr=True)
                au_rates_tables.extend(tbl["table"])
                if self.show_progress and idx % 10 == 0:
                    print(f"Table #{idx} complete. "
                          f"Total rows processed: {len(au_rates_tables)}")
                
        return au_rates_tables
    
    def convert_tables_to_dataframe(self, table):
        headers = c.AU_RATES["HEADERS"]                
        au_rates_df = self.table_to_df(table=table, headers=headers)        
        return au_rates_df
    
    def clean_and_standardise_data(self, df):        
        # convert completely blank rows to np.NaN, and drop all-NaN rows
        df = ut.convert_blank_rows_to_nan(df)
        df = ut.remove_nan_rows(df)
        
        # add a column with a standardised 10-digit comm code
        df = ut.add_standardised_comm_code(df, c.COL_HDRS["COMM_CODE"])
        
        # Fill empty / NaN cells with 'empty_cell'
        df = ut.fill_unused_cells(
            df=df, 
            non_regex_replace=["", " "], 
            non_regex_fill_str=c.CONST["EMPTY_CELL"]
        )
        
        return df
    
    def run_processor(self):
        # ingest and process the tables
        rates_table = self.process_all_tables()    
        rates_df = self.convert_tables_to_dataframe(rates_table)
        
        # clean and standardise the results
        rates_df = self.clean_and_standardise_data(rates_df)
        
        # archive old data and output results to csv, settings to txt
        dir_name = os.path.dirname(self.output_file)
        ut.auto_archive_files(dir_name, 
                               os.path.join(dir_name, c.ARCHIVE_DIR))
        rates_df.to_csv(ut.add_timestamp(self.output_file), index=False)
        ut.write_file_info(dir_name, os.path.basename(self.filepath))
        
        # return, to view head of final df
        return rates_df
    

# ------------
# AU Uses Doc
# ------------
class AuUsesDocIngest(WordDocProcessor):
    """This document contains two tables; a 'main' table, which holds
    all the Description and Usage information for each AU commodity
    code, and a short table at the end that is referred to in some of
    the usage descriptions. This module ignores the second table. 
    
    Processing undertaken:
    - Usage information is extracted from the original description and
    added as a new column
    - Commodity description is extracted from the original description
    and added as a new column
    - Completely blank rows are removed
    - Any NaN or empty cells are replaced with "empty_cell"
    - Standardised 10-digit commodity codes are added
    
    Parameters:
        - input_file (str), the file containin the reference document
        to be processed
        - output_file (str), the file to write results to
        - show_progress (bool), flag to indicate whether or not to
        print short updates to the console.
    
    NOTE:
    - The description field in the document is unusual, in that it
    includes the description for the item; the description of its 
    parents; and usage information, all in one long string. This
    requires a substantial amount of processing, which may be a source
    of error (esp if the document format changes).
    """
    def __init__(self, input_file, output_file, show_progress=True):
        super().__init__(input_file)
        self.output_file = output_file
        self.show_progress = show_progress
    
    def process_table(self):
        """Convert the first Word doc table in the list of 'tables'
        collection to a Python list.
        
        Returns:
            - list of lists, where each sub-list represents a table
            in the reference document
        
        Notes:
            - The AU Uses document only contains two tables, and the
            Uses data required is all in the first of these.
        """
        table = self.docx_to_list(self.doc.tables[0], 
                                  first_row_as_hdr=False)
        return table["table"]
    
    def convert_tables_to_dataframe(self, table):
        """Convert a python list to a pandas DataFrame, with relevant
        headers etc.
        
        Parameters:
            - table (list), data extracted from the Word doc and
            converted to a Python list
            
        Returns:
            - pandas DataFrame, containing the same table data but in
            a much more convenient format
        """
        # convert the table to a dataframe
        uses_df_full = self.table_to_df(table, c.AU_USES["HEADERS"])

        # there are a number of rows we don't need. Filter out anything
        # where the comm code isn't either entirely numeric, or where 
        # the comm code doesn't start with '0' (there are cases where
        # comm codes contain whitespace)
        uses_df = uses_df_full[
            (uses_df_full[c.COL_HDRS["COMM_CODE"]].str.isnumeric()) |
            (uses_df_full[c.COL_HDRS["COMM_CODE"]].str.startswith("0"))
        ]

        return uses_df

    def replace_trailing_text(self, df, col_name, pattern):
        idx = -1 * len(pattern)
        df[col_name] = np.where(
            df[col_name].str[idx:] == pattern, 
            df[col_name].str[:idx], 
            df[col_name]
        )
        return df
    
    def extract_usage_statements(self, df, desc_col):
        """
        """
        # set up the regex search patterns from the usage patterns. 
        # output: '(?i)<usage_pattern_1>|(?i)<usage_pattern_2>|...'
        # The '(?i)' parts make the comparison case-insensitive
        use_ptns = c.AU_USES["USAGE_PATTERNS"].keys()
        search_patterns = "|".join(f"(?i){key}" for key in use_ptns)
        
        # Split the usage terms out of the string. This will return:
        # - "number_0" will be the non-usage part of the description; 
        # - "number_1" will be the pattern found; and 
        # - "number_2" will be the remainder of the usage text (ignore)
        usage_cols = (
            df[desc_col].str.split(f"({search_patterns})",expand=True)
                        .add_prefix('number_')
        )
        
        # at least some of the usage patterns will have trailing chars; 
        # the patterns below will remove all of these.        
        for rp in c.AU_USES["REPLACE_PATTERNS"]:
            usage_cols = self.replace_trailing_text(usage_cols, "number_0", rp)

        # map the description patterns to the full description text
        usage_cols["number_1"] = usage_cols["number_1"].map(c.AU_USES["USAGE_PATTERNS"])
        
        # join the description and usage columns to the original df
        df = pd.concat([df, usage_cols[["number_0", "number_1"]]], axis=1)
        df = df.rename(columns={
            "number_0": c.AU_USES["NEW_COLS"]["FULL_DESC"], 
            "number_1": c.AU_USES["NEW_COLS"]["USAGE"]
        })

        return df
    
    def split_on_last_newline(self, df, desc_col):
        # get everything from the last newline onwards
        new_desc_cols = (df[desc_col].str.rsplit("\n", n=1, expand=True)
                                     .add_prefix("number_"))
        return new_desc_cols["number_1"].dropna()
    
    def desc_bullets(self, df, desc_col, pattern):
        """

        Parameters:
            - df (pandas DataFrame object), the dataframe to operate on
            - desc_col (str), column in df that contains description text
            - pattern (str), the pattern to split on for the *initial*
              split; e.g. "\n-" or f"\n{chr(8226)}"
        """
        # get everything from the first bullet onwards
        new_desc_cols = (df[desc_col].str.split(pattern, n=1, expand=True)
                                     .add_prefix("number_"))

        # get rid of any blank lines (would prefer not to do this, but
        # can't see an easy way to achieve the goal otherwise)
        for i in ["\n ", "\n"]:
            new_desc_cols = self.replace_trailing_text(new_desc_cols, 
                                                       "number_0", i)

        # get the last line of the 'old' description (that's now lost
        # the bulletted section)
        first_line = self.split_on_last_newline(new_desc_cols, "number_0")
        description = (first_line + pattern + new_desc_cols["number_1"])

        return description.dropna()
    
    def desc_hyphens_last_line(self, df, desc_col, pattern):
        """

        Parameters:
            - df (pandas DataFrame object), the dataframe to operate on
            - desc_col (str), column in df that contains description text
            - pattern (str), the pattern to split on for the *initial*
              split; e.g. "\n-" or f"\n{chr(8226)}"
        """
        # get everything from the first bullet onwards
        new_desc_cols = (df[desc_col].str.rsplit(pattern, n=1, expand=True)
                                     .add_prefix("number_"))
        
        # concatenate the pattern with the split column, to insert the pattern back in
        if "number_1" in new_desc_cols.columns.values:
            description = (pattern + new_desc_cols["number_1"])
        else:
            description = pd.Series(name="number_1", dtype=str)

        return description.dropna()
        
    def extract_commodity_description(self, df, desc_col, new_col_name):
        """Use each of the extraction functions in turn to extract the
        description info for a particular description pattern, removing
        the extracted rows from the overall df each time (to avoid
        clashes between extraction patterns).
        
        Parameters:
            - df (pandas DataFrame), the dataframe to operate on. It
            should contain all the rows that we want to extract
            descriptions from
            - desc_col (str), the name of the column in df that
            contains the descriptions to process. This column "SHOULD
            NOT* contain any usage information.
            - new_col_name (str), the name of the new column containing
            the processed descriptions
        
        Returns:
            - pandas Series object with the name new_col_name. This
            Series should be the same length as df 
        """
        single_hyphen_rows = df[df[desc_col] == "-"][desc_col]
        df_cut = df[~df.index.isin(single_hyphen_rows.index)]
        
        hyphens_no_spc = self.desc_hyphens_last_line(df_cut, desc_col, "\n--")
        df_cut = df_cut[~df_cut.index.isin(hyphens_no_spc.index)]
        
        hyphens_spc = self.desc_hyphens_last_line(df_cut, desc_col,"\n- -")
        df_cut = df_cut[~df_cut.index.isin(hyphens_spc.index)]
        
        multi_line_hyphens = self.desc_bullets(df_cut, desc_col, f"\n-")
        df_cut = df_cut[~df_cut.index.isin(multi_line_hyphens.index)]
        
        multi_line_bullets = self.desc_bullets(df_cut, desc_col,f"\n{chr(8226)}")
        df_cut = df_cut[~df_cut.index.isin(multi_line_bullets.index)]
        
        last_line = self.split_on_last_newline(df_cut, desc_col)
        df_cut = df_cut[~df_cut.index.isin(last_line.index)]
        
        single_line_desc = df_cut[desc_col]
        
        # concatentate all the individual Series together and return
        goods_description = (
            single_hyphen_rows.rename("number_1")
                .append(hyphens_no_spc)
                .append(hyphens_spc)
                .append(multi_line_hyphens)
                .append(multi_line_bullets)
                .append(last_line)
                .append(single_line_desc.rename("number_1"))
                .sort_index()
                .rename(new_col_name) 
        )
        return goods_description
    
    def clean_and_standardise_data(self, df):        
        # convert completely blank rows to np.NaN, and drop all-NaN rows
        df = ut.convert_blank_rows_to_nan(df)
        df = ut.remove_nan_rows(df)
        
        # add a column with a standardised 10-digit comm code
        df = ut.add_standardised_comm_code(df, c.COL_HDRS["COMM_CODE"])
        
        # Fill empty / NaN cells with 'empty_cell'
        df = ut.fill_unused_cells(
            df=df, 
            non_regex_replace=["", " "], 
            non_regex_fill_str=c.CONST["EMPTY_CELL"]
        )
        
        return df
    
    def run_processor(self):
        
        if self.show_progress:
            print("> Reading document tables...")
        uses_table = self.process_table()
        
        uses_df = self.convert_tables_to_dataframe(uses_table)
        if self.show_progress:
            print("> Extracting usage statements...")
        uses_df = self.extract_usage_statements(
            df=uses_df, 
            desc_col=c.AU_USES["HEADERS"][1]
        )
        
        if self.show_progress:
            print("> Extracting goods descriptions...")
        goods_desc = self.extract_commodity_description(
            df=uses_df, 
            desc_col=c.AU_USES["NEW_COLS"]["FULL_DESC"],
            new_col_name=c.AU_USES["NEW_COLS"]["DESC"]
        )
        uses_df = uses_df.join(goods_desc, how="inner")
        uses_df = self.clean_and_standardise_data(uses_df)
        
        # archive previous data and write new to file
        uses_df.to_csv(ut.add_timestamp(self.output_file), index=False)        
        ut.write_file_info(os.path.dirname(self.output_file), 
                            os.path.basename(self.filepath))
        
        if self.show_progress:
            print("> Processing complete.")
        
        return uses_df


#======================================================================================
# DRIVERS
#======================================================================================    
def parse_arguments():
    """Parse command-line arguments / flags, to change the default
    behaviour of the processor.
    
    Returns:
        - argparse ArgumentParser object, containing all the arguments
        provided in the namespace.
    """
    parser = argparse.ArgumentParser()
    parser.add_argument("-t", "--test",
                        help=("Set the processor to 'test' mode,"
                              " running against a small demo corpus."),
                        action="store_true")
    return parser.parse_args()


def main():
    print(f"\n{'='*50}\nSTARTING DATA INGEST\n")
    start = datetime.now()
    
    #----------------------------------
    # Parse any command-line arguments
    #----------------------------------
    args = parse_arguments()    
    if args.test:
        print(f"\n{'*'*28}\n*** RUNNING IN TEST MODE ***\n{'*'*28}\n")
    
    in_key = "INPUT_FILE_DEMO" if args.test else "INPUT_FILE"
    out_key = "OUTPUT_FILE_DEMO" if args.test else "OUTPUT_FILE"
    #----------------------------------
    
    # ------------------------------------
    # Run the AuRatesDocIngest class
    # ------------------------------------     
    print(f"Running AU Rates Extraction\n{'-'*27}")
    au_r_d = AuRatesDocIngest(c.AU_RATES[in_key], c.AU_RATES[out_key])
    df = au_r_d.run_processor()
    # ------------------------------------
    
    # ------------------------------------
    # Run the AuUsesDocIngest class
    # ------------------------------------       
    print(f"\nRunning AU Uses Extraction\n{'-'*26}")
    au_u_d = AuUsesDocIngest(c.AU_USES[in_key], c.AU_USES[out_key])
    df = au_u_d.run_processor()
    # ------------------------------------
    
    end = datetime.now()
    print(f"\nINGEST COMPLETE. Run time: {(end-start).total_seconds()}"
        f"\n{'='*50}\n")

if __name__ == "__main__":
    main()